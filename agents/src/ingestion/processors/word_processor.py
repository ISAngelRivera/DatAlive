"""
Word document processor using python-docx
"""

import logging
from typing import Dict, Any, Tuple, Union
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.oxml.exceptions import InvalidXmlError
except ImportError:
    Document = None
    InvalidXmlError = Exception

from .base_processor import BaseProcessor

logger = logging.getLogger(__name__)


class WordProcessor(BaseProcessor):
    """Word document processor"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.docx', '.doc']
        
        if not Document:
            logger.warning("python-docx not installed. Word processing will use fallback method.")
    
    async def extract_content(
        self,
        source: Union[str, Path, Dict[str, Any]],
        **kwargs
    ) -> Tuple[str, Dict[str, Any]]:
        """
        Extract content and metadata from Word document
        
        Args:
            source: Path to Word document
            **kwargs: Additional processing parameters
                - include_tables: bool = True - Include table content
                - include_headers_footers: bool = False - Include headers/footers
                - include_comments: bool = False - Include comments
                - preserve_formatting: bool = False - Preserve basic formatting
        
        Returns:
            Tuple of (content, metadata)
        """
        file_path = Path(source)
        
        if not await self.validate_source(file_path):
            raise FileNotFoundError(f"Word document not found: {file_path}")
        
        include_tables = kwargs.get('include_tables', True)
        include_headers_footers = kwargs.get('include_headers_footers', False)
        include_comments = kwargs.get('include_comments', False)
        preserve_formatting = kwargs.get('preserve_formatting', False)
        
        try:
            if Document:
                content, metadata = await self._extract_with_docx(
                    file_path, include_tables, include_headers_footers,
                    include_comments, preserve_formatting
                )
            else:
                # Fallback method
                content, metadata = await self._extract_fallback(file_path)
            
            # Add basic file metadata
            basic_metadata = self.extract_basic_metadata(file_path)
            metadata.update(basic_metadata)
            
            return self.clean_text(content), metadata
            
        except Exception as e:
            logger.error(f"Error processing Word document {file_path}: {e}")
            raise
    
    async def _extract_with_docx(
        self,
        file_path: Path,
        include_tables: bool = True,
        include_headers_footers: bool = False,
        include_comments: bool = False,
        preserve_formatting: bool = False
    ) -> Tuple[str, Dict[str, Any]]:
        """Extract content using python-docx"""
        content_parts = []
        metadata = {
            'extraction_method': 'python_docx',
            'has_tables': False,
            'has_images': False,
            'has_headers_footers': False,
            'has_comments': False
        }
        
        try:
            # Load document
            doc = Document(file_path)
            
            # Extract document metadata
            core_props = doc.core_properties
            if core_props:
                metadata.update({
                    'title': core_props.title or file_path.stem,
                    'author': core_props.author,
                    'subject': core_props.subject,
                    'keywords': core_props.keywords,
                    'comments': core_props.comments,
                    'category': core_props.category,
                    'created_at': core_props.created.isoformat() if core_props.created else None,
                    'modified_at': core_props.modified.isoformat() if core_props.modified else None,
                    'last_modified_by': core_props.last_modified_by,
                    'revision': core_props.revision
                })
            
            # Extract paragraphs
            paragraph_count = 0
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    if preserve_formatting:
                        formatted_text = self._extract_paragraph_formatting(paragraph)
                        content_parts.append(formatted_text)
                    else:
                        content_parts.append(text)
                    paragraph_count += 1
            
            metadata['paragraph_count'] = paragraph_count
            
            # Extract tables
            if include_tables and doc.tables:
                metadata['has_tables'] = True
                metadata['table_count'] = len(doc.tables)
                
                content_parts.append("\n## Tables\n")
                for i, table in enumerate(doc.tables):
                    table_content = self._extract_table_content(table, i + 1)
                    content_parts.append(table_content)
            
            # Check for images
            if self._has_images(doc):
                metadata['has_images'] = True
                content_parts.append("\n[Document contains images]")
            
            # Extract headers and footers
            if include_headers_footers:
                headers_footers = self._extract_headers_footers(doc)
                if headers_footers:
                    metadata['has_headers_footers'] = True
                    content_parts.append("\n## Headers and Footers\n")
                    content_parts.append(headers_footers)
            
            # Extract comments (if available in the document structure)
            if include_comments:
                comments = self._extract_comments(doc)
                if comments:
                    metadata['has_comments'] = True
                    content_parts.append("\n## Comments\n")
                    content_parts.append(comments)
            
        except InvalidXmlError as e:
            logger.error(f"Invalid Word document format: {e}")
            raise
        except Exception as e:
            logger.error(f"Error processing Word document: {e}")
            raise
        
        return '\n'.join(content_parts), metadata
    
    def _extract_paragraph_formatting(self, paragraph) -> str:
        """Extract paragraph with basic formatting preserved"""
        text = paragraph.text
        
        # Check paragraph style
        style_name = paragraph.style.name if paragraph.style else ""
        
        # Apply basic markdown formatting based on style
        if "Heading 1" in style_name:
            return f"# {text}"
        elif "Heading 2" in style_name:
            return f"## {text}"
        elif "Heading 3" in style_name:
            return f"### {text}"
        elif "Heading" in style_name:
            return f"#### {text}"
        else:
            return text
    
    def _extract_table_content(self, table, table_number: int) -> str:
        """Extract content from a table"""
        table_parts = [f"\n### Table {table_number}\n"]
        
        for i, row in enumerate(table.rows):
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace('\n', ' ')
                row_cells.append(cell_text)
            
            if i == 0:  # Header row
                table_parts.append('| ' + ' | '.join(row_cells) + ' |')
                table_parts.append('|' + '---|' * len(row_cells))
            else:
                table_parts.append('| ' + ' | '.join(row_cells) + ' |')
        
        return '\n'.join(table_parts) + '\n'
    
    def _has_images(self, doc) -> bool:
        """Check if document contains images"""
        try:
            # Check for inline shapes (images)
            for rel in doc.part.rels.values():
                if rel.reltype.endswith('/image'):
                    return True
            return False
        except Exception:
            return False
    
    def _extract_headers_footers(self, doc) -> str:
        """Extract headers and footers content"""
        headers_footers = []
        
        try:
            # Extract headers
            for section in doc.sections:
                # Header
                if section.header:
                    for paragraph in section.header.paragraphs:
                        text = paragraph.text.strip()
                        if text:
                            headers_footers.append(f"Header: {text}")
                
                # Footer
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        text = paragraph.text.strip()
                        if text:
                            headers_footers.append(f"Footer: {text}")
        
        except Exception as e:
            logger.warning(f"Error extracting headers/footers: {e}")
        
        return '\n'.join(headers_footers)
    
    def _extract_comments(self, doc) -> str:
        """Extract comments from document"""
        # Note: python-docx doesn't directly support comments extraction
        # This is a placeholder for future enhancement
        return ""
    
    async def _extract_fallback(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Fallback extraction method"""
        logger.warning("Using fallback Word extraction method")
        
        metadata = {
            'extraction_method': 'fallback',
            'title': file_path.stem
        }
        
        # Simple approach - just indicate Word content
        content = f"Word Document: {file_path.name}\n"
        content += "Content extraction requires python-docx library.\n"
        content += f"File size: {file_path.stat().st_size} bytes"
        
        # Try to read as plain text (limited success)
        try:
            with open(file_path, 'rb') as f:
                # Read first 1000 bytes and try to extract readable text
                data = f.read(1000)
                readable_chars = ''.join(chr(b) for b in data if 32 <= b <= 126)
                if len(readable_chars) > 50:
                    content += f"\n\nPartial content (plain text extraction):\n{readable_chars[:200]}..."
        except Exception:
            pass
        
        return content, metadata
    
    def get_document_info(self, file_path: Path) -> Dict[str, Any]:
        """Get basic document information without full extraction"""
        try:
            if not Document:
                return {'error': 'python-docx library not available'}
            
            doc = Document(file_path)
            
            info = {
                'file_size': file_path.stat().st_size,
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'has_images': self._has_images(doc)
            }
            
            # Get metadata
            core_props = doc.core_properties
            if core_props:
                info.update({
                    'title': core_props.title,
                    'author': core_props.author,
                    'created_at': core_props.created.isoformat() if core_props.created else None,
                    'modified_at': core_props.modified.isoformat() if core_props.modified else None
                })
            
            return info
            
        except Exception as e:
            logger.error(f"Error getting Word document info for {file_path}: {e}")
            return {'error': str(e)}